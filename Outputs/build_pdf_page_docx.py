from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt


ROOT = Path(__file__).resolve().parents[1]
PAGE_DIR = ROOT / "Outputs" / "main_pdf_pages_200dpi"
OUT = ROOT / "versions" / "FeatureSAGE_SCII_Com.docx"


def set_zero_spacing(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1


def main():
    images = sorted(PAGE_DIR.glob("main-*.png"), key=lambda p: int(p.stem.split("-")[-1]))
    if not images:
        raise SystemExit(f"No page images found in {PAGE_DIR}")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)

    for i, image in enumerate(images):
        paragraph = doc.add_paragraph()
        set_zero_spacing(paragraph)
        if i:
            paragraph.paragraph_format.page_break_before = True
        run = paragraph.add_run()
        run.add_picture(str(image), width=Mm(207))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
